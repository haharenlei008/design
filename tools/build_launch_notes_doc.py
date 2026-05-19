from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "AI投影灯功能上线说明_v1.0.0.docx"


COLORS = {
    "ink": "1B1834",
    "muted": "6B6884",
    "purple": "342A63",
    "violet": "6D4CF5",
    "gold": "F0B84D",
    "blue": "2F7DFF",
    "fill": "F6F3FF",
    "fill2": "FFF7E6",
    "border": "D9D2F0",
    "white": "FFFFFF",
}


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "D9D2F0", size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa: int = 9360) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    font = run.font
    font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size is not None:
        font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if color is not None:
        font.color.rgb = rgb(color)


def add_run(paragraph, text: str, size: float | None = None, bold: bool | None = None, color: str | None = None):
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return run


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_font(run, size={1: 17, 2: 14, 3: 12}.get(level, 12), bold=True, color=COLORS["purple"] if level == 1 else COLORS["ink"])
    return paragraph


def add_body(doc: Document, text: str, color: str = "ink", after: int = 7):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.2
    add_run(paragraph, text, size=10.5, color=COLORS[color])
    return paragraph


def add_bullet(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.18
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.12)
    add_run(paragraph, text, size=10.2, color=COLORS["ink"])
    return paragraph


def add_number(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.18
    add_run(paragraph, text, size=10.2, color=COLORS["ink"])
    return paragraph


def add_callout(doc: Document, title: str, body: str, fill: str = "F6F3FF"):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, COLORS["border"])
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)
    p = cell.paragraphs[0]
    add_run(p, title, size=10.5, bold=True, color=COLORS["purple"])
    p.add_run("\n")
    add_run(p, body, size=10.2, color=COLORS["ink"])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_simple_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table)
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_shading(header_cells[index], "EFE9FF")
        set_cell_border(header_cells[index], COLORS["border"])
        paragraph = header_cells[index].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(paragraph, header, size=9.5, bold=True, color=COLORS["purple"])
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_border(cells[index], "E6E0F5", "6")
            cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            add_run(paragraph, value, size=9.2, color=COLORS["ink"])
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_image_gallery(doc: Document, items: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_width(table)
    for idx in range(0, len(items), 2):
        row = table.add_row().cells
        for col in range(2):
            item_index = idx + col
            cell = row[col]
            set_cell_shading(cell, "FBFAFF")
            set_cell_border(cell, "E6E0F5", "6")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if item_index >= len(items):
                continue
            title, file_name = items[item_index]
            image_path = ROOT / file_name
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, title, size=9.2, bold=True, color=COLORS["purple"])
            if image_path.exists():
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(image_path), width=Inches(2.55))
            else:
                add_run(cell.add_paragraph(), f"缺少图片：{file_name}", size=9, color="9B1C1C")
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(COLORS["ink"])
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.2
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.color.rgb = rgb(COLORS["purple"])
    return doc


def build() -> None:
    doc = setup_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    add_run(title, "AI 宝宝投影灯", size=23, bold=True, color=COLORS["purple"])
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    add_run(subtitle, "v1.0.0 功能上线说明", size=16, bold=True, color=COLORS["gold"])

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(meta, "适用范围：欢迎页、登录页、设备绑定流程、设备首页在线态与离线态  |  编制日期：2026 年 5 月", size=9.5, color=COLORS["muted"])

    add_callout(
        doc,
        "上线目标",
        "本次上线围绕“让家长快速进入陪伴体验”展开，完整覆盖首次打开、登录、设备绑定、在线控制、离线查看与核心内容生成链路。页面整体强调温暖、安心、童趣和低学习成本。",
        "FFF7E6",
    )

    add_heading(doc, "1. 本次上线范围", 1)
    add_simple_table(
        doc,
        ["模块", "上线内容", "用户价值"],
        [
            ["欢迎页", "品牌初始展示、产品氛围引导、进入登录入口。", "建立产品第一印象，降低首次打开的陌生感。"],
            ["登录页", "账号登录入口、基础身份确认、继续使用路径。", "完成用户身份闭环，为后续设备与宝宝信息承接做准备。"],
            ["设备未绑定页", "空状态说明、绑定设备引导、无设备时的路径提示。", "让用户清楚知道下一步是添加或绑定投影灯。"],
            ["设备绑定页", "设备绑定结果与进入首页路径。", "完成设备关系建立，让用户进入正式控制体验。"],
            ["设备首页", "模式切换、随机播放、AI 创作、语音、亮度、夜灯、停止播放、设备切换、侧栏。", "集中承载日常陪伴、播放和设备控制。"],
            ["设备离线页", "展示设备不在线状态，保留模式浏览与账户侧栏，禁用设备操作。", "离线场景下仍可查看内容结构，并避免误操作。"],
        ],
        widths=[1.28, 3.05, 2.25],
    )

    add_heading(doc, "2. 页面图文总览", 1)
    add_body(doc, "以下图片为本次上线的关键页面示意，覆盖从首次进入到设备绑定后的主要体验路径。", "muted")
    add_image_gallery(
        doc,
        [
            ("欢迎页", "welcome_page_ui_assets.png"),
            ("登录页", "login_page_ui_assets.png"),
            ("设备未绑定页", "设备未绑定.png"),
            ("设备绑定页", "设备绑定页.png"),
            ("首页：故事模式", "故事模式-设计图.png"),
            ("首页：儿歌模式", "儿歌模式-设计图.png"),
            ("首页：睡眠模式", "睡眠模式-设计图.png"),
            ("首页：亲子模式", "亲子模式-设计图.png"),
        ],
    )

    add_heading(doc, "3. 首页核心功能说明", 1)
    add_heading(doc, "3.1 设备顶部区", 2)
    add_bullet(doc, "展示当前设备名称与在线状态，设备名称支持下拉切换。")
    add_bullet(doc, "设备列表中可切换已绑定设备；“添加设备”作为静态入口展示，不跳转新页面。")
    add_bullet(doc, "右上角头像进入侧栏，侧栏保留头像、宝宝昵称、陪伴文案、登录账号和设备状态。")

    add_heading(doc, "3.2 模式切换与动态氛围", 2)
    add_bullet(doc, "支持故事、儿歌、睡眠、亲子四种模式。切换后同步更新背景、标题、副标题、随机播放文案与 AI 创作提示。")
    add_bullet(doc, "模式点击有轻微下压与选中扫光反馈，背景采用交叉淡入淡出。")
    add_bullet(doc, "不同模式具备专属背景动效：故事星星闪动、儿歌音符轻跳、睡眠月光与 Z 符号呼吸、亲子爱心柔光上浮。")

    add_heading(doc, "3.3 随机播放", 2)
    add_bullet(doc, "根据当前模式推荐一个主题，弹窗中可“换一个”并确认发送到设备播放。")
    add_bullet(doc, "主题名称按最长约 20 个字处理，避免标题过大或溢出。")
    add_bullet(doc, "发送过程展示“发送中”和“发送成功”，成功后首页播放卡片同步更新。")

    add_heading(doc, "3.4 AI 创作", 2)
    add_bullet(doc, "用户输入生成要求后进入模拟生成状态，当前原型约 5 秒完成。")
    add_bullet(doc, "生成完成后页面只保留“重新生成”和“发送给设备”两个关键操作，减少中间音频流信息干扰。")
    add_bullet(doc, "发送给设备后进入发送中与成功反馈，成功后首页播放卡片更新为 AI 生成内容。")
    add_bullet(doc, "生成中阶段支持取消生成；结果阶段支持关闭弹窗或重新生成。")

    add_heading(doc, "3.5 语音留言", 2)
    add_bullet(doc, "底部中间“语音”为核心入口，长按 1.5 秒后进入录音展示。")
    add_bullet(doc, "松开结束录音，展示已录制音频，可试听后发送到设备播放。")
    add_bullet(doc, "语音录音预览支持关闭弹窗；去掉重录按钮，流程更轻。")

    add_heading(doc, "3.6 设备控制", 2)
    add_bullet(doc, "底部保留投影亮度、语音、夜灯三项，语音置中并强化视觉。")
    add_bullet(doc, "投影亮度点击打开滑杆面板，滑动实时展示百分比。")
    add_bullet(doc, "夜灯点击切换开关状态。")
    add_bullet(doc, "页面提供停止播放功能，点击后向设备发送停止播放指令，并在页面中间提示“已停止播放”。")

    add_heading(doc, "4. 侧栏功能说明", 1)
    add_body(doc, "右上角区域是本次首页的重要收纳入口，主要包含“设备选择”和“头像侧栏”两类能力：设备相关操作靠近设备名称，账户与功能中心收纳在头像入口中。")
    add_simple_table(
        doc,
        ["入口/区域", "内容", "说明"],
        [
            ["设备名称下拉", "当前设备、在线状态、设备列表、添加设备", "点击设备名称可展开设备列表；选择设备只更新当前展示；添加设备仅做入口反馈，不跳转。"],
            ["头像入口", "右上角宝宝头像按钮", "点击后打开右侧功能侧栏；打开时背景柔和遮罩，支持关闭按钮、遮罩和 Escape 关闭。"],
            ["资料区", "头像、宝宝昵称、陪伴文案、登录账号、设备状态", "用于承载账户和设备的轻量信息，让用户确认当前登录账号与设备在线状态。"],
            ["主功能", "播放记录、设备设置、陪伴文案", "点击后在侧栏内给出反馈；离线时设备设置提示连接后可调整。"],
            ["更多设置", "帮助与支持、关于系统、其他", "聚合在线客服、常见问题、使用说明、反馈建议、版本、隐私、服务协议、加盟合作、在线购买。"],
            ["交互反馈", "功能反馈提示", "侧栏不新增页面、不跳转；点击入口后只展示对应说明或状态提示。"],
        ],
        widths=[1.25, 2.6, 2.75],
    )

    add_heading(doc, "4.1 右上角功能清单", 2)
    add_bullet(doc, "头像侧栏：作为账户与功能中心入口，承载播放记录、设备设置、陪伴文案和更多设置。")
    add_bullet(doc, "设备下拉：用户可在首页顶部快速切换已绑定设备，并看到设备在线/不在线状态。")
    add_bullet(doc, "添加设备入口：在设备下拉中展示，当前原型只显示入口反馈，不进入配网页面。")
    add_bullet(doc, "离线态适配：离线页仍可打开头像侧栏和设备信息展示，但设备设置会提示连接后可调整。")

    add_heading(doc, "5. 离线页说明", 1)
    add_body(doc, "离线页整体内容与首页保持一致，但顶部设备状态显示“设备不在线”。离线状态下模式切换、背景动效和头像侧栏仍可用，设备与内容操作保持禁用。")
    add_bullet(doc, "禁用项包括随机播放、AI 立即生成、亮度、夜灯、语音、停止播放等设备相关操作。")
    add_bullet(doc, "模式卡片可正常切换，帮助用户预览不同模式的内容和氛围。")
    add_bullet(doc, "侧栏保留账户相关入口，设备设置在离线状态给出限制提示。")

    add_heading(doc, "6. 主要用户路径", 1)
    add_number(doc, "首次打开：欢迎页展示产品氛围，用户进入登录页。")
    add_number(doc, "账号确认：登录页完成身份确认，进入设备状态判断。")
    add_number(doc, "未绑定设备：展示设备未绑定页，引导进入绑定流程。")
    add_number(doc, "完成绑定：设备绑定页提示绑定成功，进入设备首页。")
    add_number(doc, "日常使用：首页选择模式，可随机播放、AI 创作、语音发送、调节亮度和夜灯。")
    add_number(doc, "设备离线：进入离线页，查看模式和账户信息，设备操作被禁用。")

    add_heading(doc, "7. 测试验收清单", 1)
    add_simple_table(
        doc,
        ["检查项", "验收标准"],
        [
            ["页面入口", "欢迎页、登录页、设备未绑定页、设备绑定页、首页、离线页均可打开，无明显布局错位。"],
            ["设备状态", "在线页显示设备在线，离线页显示设备不在线，离线操作按钮不可触发。"],
            ["模式切换", "四种模式切换后背景、文案、播放推荐、AI 提示与动效同步变化。"],
            ["随机播放", "可换一个主题，确认后有发送中和成功反馈，首页播放卡片更新。"],
            ["AI 创作", "输入为空有提示；输入后约 5 秒生成完成；可重新生成或发送给设备。"],
            ["语音", "长按 1.5 秒进入录音，松开完成；可试听并发送；关闭预览后页面恢复正常。"],
            ["设备控制", "亮度可调节，夜灯可切换，停止播放点击后页面中间提示已停止播放。"],
            ["侧栏", "头像可打开侧栏，主功能和更多设置有反馈，关闭按钮/遮罩/Escape 均可关闭。"],
        ],
        widths=[1.55, 4.95],
    )

    add_heading(doc, "8. 上线注意事项", 1)
    add_callout(
        doc,
        "原型边界",
        "当前 AI 创作、随机播放发送、语音发送、添加设备等均为静态原型或模拟反馈，不接真实后端、不上传真实设备。后续接入后端时，可保留现有 UI 状态机，只替换请求和设备通信层。",
        "F6F3FF",
    )
    add_bullet(doc, "长文本需要持续压测，包括 AI 生成标题、随机主题名称、登录账号和设备名称。")
    add_bullet(doc, "设备离线场景要避免触发麦克风、生成请求、播放发送、亮度和夜灯控制。")
    add_bullet(doc, "移动端视口需重点检查 595px 与 682px 宽度，确保底部栏、弹窗和侧栏不遮挡核心内容。")
    add_bullet(doc, "真实后端上线前，需要补充接口错误、超时、重试、设备不可达和音频生成失败等异常态。")

    add_heading(doc, "9. 后续建议", 1)
    add_bullet(doc, "接入真实设备列表与设备绑定状态，替换当前静态设备下拉数据。")
    add_bullet(doc, "AI 创作接入真实音频生成接口后，补充生成失败、排队中、内容审核未通过等状态。")
    add_bullet(doc, "播放记录页、设备设置页和陪伴文案页可在下一阶段从侧栏主功能中逐步落地。")
    add_bullet(doc, "为语音留言补充真实录音权限处理、上传进度和设备播放回执。")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "AI 宝宝投影灯 v1.0.0 功能上线说明", size=8.5, color=COLORS["muted"])

    doc.save(OUT)


if __name__ == "__main__":
    build()

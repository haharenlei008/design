from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "launch_notes_assets" / "weekly_2026_05_22"
OUT = ROOT / "AI投影灯（内测版）本周新增功能上线说明_2026-05-22.docx"

COLORS = {
    "ink": "1B1834",
    "muted": "625F78",
    "heading": "2E2472",
    "accent": "6D4CF5",
    "gold": "EBAA35",
    "blue": "2563EB",
    "green": "14805F",
    "red": "B42318",
    "fill": "F6F3FF",
    "fill_gold": "FFF7E6",
    "fill_blue": "EFF6FF",
    "border": "D9D2F0",
    "soft_border": "E9E2F8",
}


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    font = run.font
    font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size is not None:
        font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if color is not None:
        font.color.rgb = rgb(color)


def add_run(paragraph, text: str, size: float | None = None, bold: bool | None = None, color: str | None = None):
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color or COLORS["ink"])
    return run


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = COLORS["soft_border"], size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_in: list[float], indent_dxa: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)

    widths_dxa = [int(round(width * 1440)) for width in widths_in]
    for width_dxa in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width_dxa))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_in[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt({1: 18, 2: 14, 3: 10}.get(level, 10))
    paragraph.paragraph_format.space_after = Pt({1: 10, 2: 7, 3: 5}.get(level, 5))
    add_run(paragraph, text, size={1: 16, 2: 13, 3: 12}.get(level, 12), bold=True, color=COLORS["heading"])
    return paragraph


def add_body(doc: Document, text: str, color: str = "ink", after: int = 6):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    add_run(paragraph, text, size=10.6, color=COLORS[color])
    return paragraph


def add_bullet(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    add_run(paragraph, text, size=10.4, color=COLORS["ink"])
    return paragraph


def add_callout(doc: Document, title: str, body: str, fill: str = COLORS["fill"]):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, COLORS["border"])
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    add_run(paragraph, title, size=10.8, bold=True, color=COLORS["heading"])
    paragraph.add_run("\n")
    add_run(paragraph, body, size=10.4, color=COLORS["ink"])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, "E8EEF5")
        set_cell_border(cell, COLORS["border"])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        add_run(paragraph, header, size=9.4, bold=True, color=COLORS["heading"])

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            set_cell_border(cell, COLORS["soft_border"], "6")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.18
            if index == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(paragraph, value, size=9.1, bold=True, color=COLORS["heading"])
            else:
                add_run(paragraph, value, size=9.1, color=COLORS["ink"])
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_feature_image(doc: Document, title: str, image_name: str, caption: str) -> None:
    image_path = ASSET_DIR / image_name
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FBFAFF")
    set_cell_border(cell, COLORS["soft_border"], "6")

    title_p = cell.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    add_run(title_p, title, size=9.5, bold=True, color=COLORS["heading"])

    if image_path.exists():
        img_p = cell.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.paragraph_format.space_after = Pt(4)
        img_p.add_run().add_picture(str(image_path), width=Inches(3.05))
    else:
        add_run(cell.add_paragraph(), f"缺少截图：{image_name}", size=9, color=COLORS["red"])

    caption_p = cell.add_paragraph()
    caption_p.paragraph_format.space_after = Pt(0)
    caption_p.paragraph_format.line_spacing = 1.15
    add_run(caption_p, caption, size=8.8, color=COLORS["muted"])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(COLORS["ink"])
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ("Heading 1", "Heading 2", "Heading 3", "List Bullet"):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.color.rgb = rgb(COLORS["heading"])

    return doc


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    add_run(title, "AI 宝宝投影灯", size=25, bold=True, color=COLORS["heading"])

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    add_run(subtitle, "内测版 · 本周新增功能上线说明", size=15, bold=True, color=COLORS["gold"])

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(14)
    add_run(
        meta,
        "统计口径：仅覆盖 2026-05-19 至 2026-05-20 git 提交新增/优化功能；不重复描述此前已实现的欢迎、登录、绑定、基础首页和基础模式切换。",
        size=9.4,
        color=COLORS["muted"],
    )

    add_callout(
        doc,
        "本周上线结论",
        "本周代码增量集中在首页体验闭环：AI 生成从“等待结果”升级为“文案可审 + 音频流可听 + 完成可发”；设备控制从单点亮度升级为“模式自动匹配灯光和旋转 + 可手动微调”；同时补齐随机试听、离线禁用和模式卡片反馈。",
        COLORS["fill_gold"],
    )


def build() -> None:
    doc = setup_document()
    add_cover(doc)

    add_heading(doc, "1. Git 提交范围", 1)
    add_table(
        doc,
        ["提交", "日期", "本次纳入的功能增量"],
        [
            ["6f14ff8", "2026-05-19", "AI 立即生成改为音频流式接收；新增 AI 试听片段、完整试听、发送设备；随机播放弹窗新增试听。"],
            ["8061be8", "2026-05-19", "新增灯光/旋转场景联动、模式默认设备效果、长亮三色亮度、场景效果开关、离线禁用补齐。"],
            ["26a91c2", "2026-05-20", "新增 AI 文案预览流程：生成文本、编辑文案、重新生成文本、确认后再生成音频。"],
            ["1b0cd7e", "2026-05-20", "优化模式卡片文字可读性、两行截断、选中态遮罩与重复点击反馈。"],
        ],
        widths=[0.85, 1.05, 4.6],
    )
    add_body(doc, "说明：合并提交、README-only 提交和上周已完成的基础能力不作为独立功能展开。", "muted")

    add_heading(doc, "2. 新增功能图文说明", 1)

    add_heading(doc, "2.1 AI 文案预览：确认文本后再生成音频", 2)
    add_feature_image(
        doc,
        "图 1：AI 文案预览弹窗",
        "01-ai-transcript-preview.png",
        "提交 26a91c2 新增 text-preview 面板。用户输入生成要求后，先看到音频对应文本，可返回修改、重新生成文本，或确认生成音频。",
    )
    add_bullet(doc, "新增对应文本输入框、空文本校验和编辑态错误提示。")
    add_bullet(doc, "新增文本版本切换逻辑，重新生成文本时不会直接覆盖已确认音频。")
    add_bullet(doc, "音频生成前置为“家长先审文案”，减少生成内容直接进入播放链路的不确定性。")

    add_heading(doc, "2.2 AI 音频流：流式接收、片段试听、完成后发送", 2)
    add_feature_image(
        doc,
        "图 2：AI 音频流接收中",
        "02-ai-audio-stream-preview.png",
        "提交 6f14ff8 新增音频流接收状态。页面显示已接收时长、总时长、进度条和流状态；收到片段后即可试听，流完成后解锁发送给设备。",
    )
    add_bullet(doc, "新增 AI 音频卡片、进度条、当前时长/总时长和“试听片段/播放试听”状态切换。")
    add_bullet(doc, "新增音频 Blob 模拟与播放控件，流式接收中可加载已收到片段。")
    add_bullet(doc, "发送前校验音频流是否完成；发送成功后首页播放卡片更新为 AI 内容。")

    add_heading(doc, "2.3 随机播放：新增主题试听", 2)
    add_feature_image(
        doc,
        "图 3：随机播放主题试听",
        "03-random-preview.png",
        "提交 6f14ff8 在随机播放弹窗中加入试听按钮和隐藏 audio 元素。用户可以先试听随机主题片段，再决定是否发送到设备播放。",
    )
    add_bullet(doc, "新增随机试听音频的生成、播放、暂停和关闭弹窗后的资源释放。")
    add_bullet(doc, "换一个主题时会清理上一段试听，避免旧音频继续播放。")

    add_heading(doc, "2.4 灯光控制：模式联动和长亮三色亮度", 2)
    add_feature_image(
        doc,
        "图 4：灯光设置面板",
        "04-light-settings.png",
        "提交 8061be8 新增灯光设置面板。支持渐变、跑马、呼吸、长亮；长亮模式下可分别调节红光、黄光、蓝光亮度。",
    )
    add_table(
        doc,
        ["能力", "本周新增点"],
        [
            ["模式默认值", "故事/儿歌/睡眠/亲子分别带默认灯光效果、速度、亮度和提示文案。"],
            ["手动微调", "灯光效果、执行速度、整体亮度和长亮三色通道都可在面板内调整。"],
            ["设备映射", "新增 prototype payload：RGB 效果、通道亮度、关闭灯光等操作都有对应映射。"],
            ["快捷状态", "底部灯光按钮会同步显示当前效果，例如渐变灯、长亮、已关闭。"],
        ],
        widths=[1.35, 5.15],
    )

    add_heading(doc, "2.5 旋转控制：三档速度和单独停转", 2)
    add_feature_image(
        doc,
        "图 5：旋转设置面板",
        "05-rotation-settings.png",
        "提交 8061be8 新增旋转设置面板。用户可选择低速、中速、高速，也可以只停止投影旋转，不影响灯光当前状态。",
    )
    add_bullet(doc, "模式切换会自动匹配旋转速度：故事低速、儿歌高速、睡眠低速、亲子中速。")
    add_bullet(doc, "新增 motor.start、motor.adjust_speed、motor.stop 的原型 payload 映射。")
    add_bullet(doc, "旋转快捷按钮与面板状态同步，操作后用 toast 给出明确反馈。")

    add_heading(doc, "2.6 模式卡片：可读性和重复点击反馈", 2)
    add_feature_image(
        doc,
        "图 6：模式卡片反馈",
        "06-mode-card-feedback.png",
        "提交 1b0cd7e 优化模式卡片标题可读性，并在重复点击当前模式时给出“当前已是该模式”的反馈，避免用户误以为点击无响应。",
    )
    add_bullet(doc, "卡片标题增加遮罩、文字阴影、最小字号和两行截断策略。")
    add_bullet(doc, "选中态和按下态反馈增强，重复点击不触发模式重置，只触发轻反馈。")
    add_bullet(doc, "README 补充模式卡片文案与本地化长度规则。")

    add_heading(doc, "2.7 离线态：设备控制禁用补齐", 2)
    add_feature_image(
        doc,
        "图 7：离线态禁用设备操作",
        "07-offline-disabled-controls.png",
        "提交 8061be8 同步更新离线页结构。离线状态下灯光、旋转等设备操作保持禁用，并提示设备在线后才能调整。",
    )
    add_bullet(doc, "离线页保留模式浏览，但禁用灯光、旋转、AI 生成、随机播放等需要在线设备的操作。")
    add_bullet(doc, "离线页也包含灯光/旋转面板结构，所有选项 disabled，确保在线页和离线页结构一致。")

    add_heading(doc, "3. 验收重点", 1)
    add_table(
        doc,
        ["功能", "验收标准"],
        [
            ["AI 文案预览", "输入需求后先进入对应文本预览；可编辑、返回修改、重新生成；空文本不能确认。"],
            ["AI 音频流", "确认文案后进入准备和流式接收；片段可试听；流完成后才能发送设备。"],
            ["随机试听", "随机播放弹窗可试听；切换主题或关闭弹窗会停止并释放上一段试听。"],
            ["灯光控制", "四种灯效、速度、整体亮度、长亮三色亮度均可调，底部状态同步。"],
            ["旋转控制", "低/中/高速和停止旋转可切换，模式默认旋转值正确应用。"],
            ["离线态", "离线页设备控制均 disabled，不触发生成、播放发送、灯光或旋转操作。"],
            ["卡片反馈", "长标题不遮挡卡片图案；重复点击当前模式有 toast 反馈，不重复重置状态。"],
        ],
        widths=[1.25, 5.25],
    )

    add_heading(doc, "4. 不在本说明展开的既有功能", 1)
    add_body(
        doc,
        "欢迎页、登录页、设备绑定页、基础首页布局、基础模式切换、侧栏入口、语音留言基础录制等不是本周新增重点，本说明不重复描述；如需全量功能手册，可另出完整版。",
        "muted",
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "AI 宝宝投影灯内测版 · 本周新增功能上线说明 · 2026-05-22", size=8.5, color=COLORS["muted"])

    doc.save(OUT)


if __name__ == "__main__":
    build()
